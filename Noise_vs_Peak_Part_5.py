import os
import h5py
import numpy as np
import tkinter as tk
from tkinter import ttk, filedialog, messagebox
import matplotlib.pyplot as plt
from matplotlib.backends.backend_tkagg import FigureCanvasTkAgg
from scipy.optimize import curve_fit
from docx import Document


# Gaussian Definition
def gaussian(x, A, mu, sigma, C):
    return A * np.exp(-(x - mu) ** 2 / (2 * sigma ** 2)) + C


class H5DataAnalyzerApp:
    def __init__(self, root):
        self.root = root
        self.root.title("HDF5 Continuous Pulse & Peak Analyzer")
        self.root.geometry("1300x900")
        self.root.minsize(1100, 750)

        # File and Analysis State Variables
        self.h5_path = None
        self.datasets_dict = {}  # maps display path to structural dataset path
        self.is_running = False

        # Iteration Variables for Continuous Processing
        self.columns_to_process = []
        self.current_col_index = 0
        self.all_peak_counts = []
        self.peak_results = []
        self.doc_summary_table_data = []
        self.active_dataset_name = ""

        self.setup_ui()

    def setup_ui(self):
        # ------------------ LEFT SIDEBAR (Controls & Tree) ------------------
        left_panel = ttk.Frame(self.root, padding=10)
        left_panel.pack(side=tk.LEFT, fill=tk.Y, expand=False)

        # File Selection Button
        file_frame = ttk.LabelFrame(left_panel, text=" HDF5 Input Storage ", padding=5)
        file_frame.pack(fill=tk.X, pady=5)

        btn_browse = ttk.Button(file_frame, text="Select HDF5 File", command=self.load_hdf5_file)
        btn_browse.pack(fill=tk.X, pady=2)

        self.lbl_filename = ttk.Label(file_frame, text="No file selected", wraplength=250, foreground="gray")
        self.lbl_filename.pack(fill=tk.X, pady=2)

        # Treeview to display inner layout of datasets
        tree_frame = ttk.LabelFrame(left_panel, text=" Select Internal Dataset ", padding=5)
        tree_frame.pack(fill=tk.BOTH, expand=True, pady=5)

        self.tree = ttk.Treeview(tree_frame, columns=("Shape", "Dtype"), show="tree headings", selectmode="browse")
        self.tree.heading("#0", text="Dataset Hierarchy")
        self.tree.heading("Shape", text="Shape")
        self.tree.heading("Dtype", text="Data Type")
        self.tree.column("#0", width=160, anchor="w")
        self.tree.column("Shape", width=90, anchor="center")
        self.tree.column("Dtype", width=60, anchor="center")

        tree_scroll = ttk.Scrollbar(tree_frame, orient=tk.VERTICAL, command=self.tree.yview)
        self.tree.configure(yscrollcommand=tree_scroll.set)

        self.tree.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
        tree_scroll.pack(side=tk.RIGHT, fill=tk.Y)
        self.tree.bind("<<TreeviewSelect>>", self.on_dataset_selected)

        # Shape Data Configuration Parameters
        config_frame = ttk.LabelFrame(left_panel, text=" Row / Column Configuration ", padding=5)
        config_frame.pack(fill=tk.X, pady=5)

        # Row ranges
        row_lbl_frame = ttk.Frame(config_frame)
        row_lbl_frame.pack(fill=tk.X, pady=2)
        ttk.Label(row_lbl_frame, text="Row Bounds (Start:End):").pack(side=tk.LEFT)
        self.ent_row_start = ttk.Entry(row_lbl_frame, width=8)
        self.ent_row_start.pack(side=tk.LEFT, padx=2)
        self.ent_row_start.insert(0, "1")
        ttk.Label(row_lbl_frame, text=" : ").pack(side=tk.LEFT)
        self.ent_row_end = ttk.Entry(row_lbl_frame, width=8)
        self.ent_row_end.pack(side=tk.LEFT, padx=2)
        self.ent_row_end.insert(0, "200003")

        # Column ranges
        col_lbl_frame = ttk.Frame(config_frame)
        col_lbl_frame.pack(fill=tk.X, pady=2)
        ttk.Label(col_lbl_frame, text="Col Bounds (Start:End):").pack(side=tk.LEFT)
        self.ent_col_start = ttk.Entry(col_lbl_frame, width=8)
        self.ent_col_start.pack(side=tk.LEFT, padx=2)
        self.ent_col_start.insert(0, "0")
        ttk.Label(col_lbl_frame, text=" : ").pack(side=tk.LEFT)
        self.ent_col_end = ttk.Entry(col_lbl_frame, width=8)
        self.ent_col_end.pack(side=tk.LEFT, padx=2)
        self.ent_col_end.insert(0, "5")

        # Pre-scan Baseline Statistics Display Label
        self.lbl_pre_scan_stats = ttk.Label(config_frame, text="Baseline: Select a channel dataset...",
                                            font=("Segoe UI", 9, "bold"), foreground="#0056b3")
        self.lbl_pre_scan_stats.pack(fill=tk.X, pady=(6, 2))

        # Central Region Threshold Customizer Panel (UPDATED WITH SECOND SLIDER)
        central_frame = ttk.LabelFrame(left_panel, text=" Central Region Threshold Customizer ", padding=5)
        central_frame.pack(fill=tk.X, pady=5)

        central_bounds_frame = ttk.Frame(central_frame)
        central_bounds_frame.pack(fill=tk.X, pady=2)
        ttk.Label(central_bounds_frame, text="Central Range (Indices):").pack(side=tk.LEFT)
        self.ent_central_start = ttk.Entry(central_bounds_frame, width=8)
        self.ent_central_start.pack(side=tk.LEFT, padx=2)
        self.ent_central_start.insert(0, "102000")
        ttk.Label(central_bounds_frame, text=" : ").pack(side=tk.LEFT)
        self.ent_central_end = ttk.Entry(central_bounds_frame, width=8)
        self.ent_central_end.pack(side=tk.LEFT, padx=2)
        self.ent_central_end.insert(0, "120000")

        # SECOND SLIDER SETTINGS
        self.central_slider_val = tk.DoubleVar(value=7.0)

        ttk.Label(central_frame, text="Central Multiplier Factor (n_c):").pack(anchor="w", pady=(5, 0))
        central_slider_row = ttk.Frame(central_frame)
        central_slider_row.pack(fill=tk.X, pady=2)

        self.scale_central_mult = ttk.Scale(central_slider_row, from_=1.0, to=15.0, variable=self.central_slider_val,
                                            orient=tk.HORIZONTAL)
        self.scale_central_mult.pack(side=tk.LEFT, fill=tk.X, expand=True, padx=(0, 5))

        self.ent_central_mult = ttk.Entry(central_slider_row, textvariable=self.central_slider_val, width=6)
        self.ent_central_mult.pack(side=tk.RIGHT)
        self.central_slider_val.trace_add("write", self.on_central_multiplier_changed)

        self.lbl_central_formula = ttk.Label(central_frame, text="Central Threshold = μ - 7.00σ",
                                             font=("Consolas", 10, "bold"), foreground="#007bff")
        self.lbl_central_formula.pack(anchor="w", pady=(4, 2))

        # FIRST SLIDER (Global Multiplier)
        slider_frame = ttk.LabelFrame(left_panel, text=" Global Multiplier Factor (n) ", padding=5)
        slider_frame.pack(fill=tk.X, pady=5)

        self.slider_val = tk.DoubleVar(value=5.0)

        controls_row = ttk.Frame(slider_frame)
        controls_row.pack(fill=tk.X, pady=2)

        self.scale_mult = ttk.Scale(controls_row, from_=1.0, to=15.0, variable=self.slider_val, orient=tk.HORIZONTAL)
        self.scale_mult.pack(side=tk.LEFT, fill=tk.X, expand=True, padx=(0, 5))

        self.ent_mult = ttk.Entry(controls_row, textvariable=self.slider_val, width=6)
        self.ent_mult.pack(side=tk.RIGHT)
        self.slider_val.trace_add("write", self.on_multiplier_changed)

        self.lbl_slider = ttk.Label(slider_frame, text="Current Multiplier: 5.00")
        self.lbl_slider.pack(anchor="w", pady=(2, 0))

        self.lbl_formula = ttk.Label(slider_frame, text="Global Threshold = μ - 5.00σ",
                                     font=("Consolas", 10, "bold"), foreground="#d9534f")
        self.lbl_formula.pack(anchor="w", pady=(4, 2))

        # Action Control Buttons
        action_frame = ttk.Frame(left_panel, padding=5)
        action_frame.pack(fill=tk.X, pady=5)

        self.btn_play = ttk.Button(action_frame, text="▶ Run Continuous", command=self.start_processing)
        self.btn_play.pack(side=tk.LEFT, fill=tk.X, expand=True, padx=2)

        self.btn_stop = ttk.Button(action_frame, text="⏸ Pause / Stop", command=self.stop_processing, state=tk.DISABLED)
        self.btn_stop.pack(side=tk.RIGHT, fill=tk.X, expand=True, padx=2)

        # ------------------ RIGHT DISPLAY PANEL (Plots & Metrics) ------------------
        right_panel = ttk.Frame(self.root, padding=10)
        right_panel.pack(side=tk.RIGHT, fill=tk.BOTH, expand=True)

        plot_frame = ttk.LabelFrame(right_panel, text=" Continuous Waveform and Fitted Gaussian Output Window ")
        plot_frame.pack(fill=tk.BOTH, expand=True)

        self.fig, self.ax = plt.subplots(figsize=(6, 4))
        self.canvas = FigureCanvasTkAgg(self.fig, master=plot_frame)
        self.canvas.get_tk_widget().pack(fill=tk.BOTH, expand=True)

        console_frame = ttk.LabelFrame(right_panel, text=" Live Computation Metrics Summary Tracker ", height=180)
        console_frame.pack(fill=tk.X, side=tk.BOTTOM, pady=5)
        console_frame.pack_propagate(False)

        self.txt_console = tk.Text(console_frame, wrap=tk.WORD, height=8, background="#222", foreground="#0f0",
                                   font=("Consolas", 9))
        console_scroll = ttk.Scrollbar(console_frame, orient=tk.VERTICAL, command=self.txt_console.yview)
        self.txt_console.configure(yscrollcommand=console_scroll.set)

        self.txt_console.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
        console_scroll.pack(side=tk.RIGHT, fill=tk.Y)

    def on_multiplier_changed(self, *args):
        try:
            val = self.slider_val.get()
            self.lbl_slider.config(text=f"Current Multiplier: {val:.2f}")
            self.lbl_formula.config(text=f"Global Threshold = μ - {val:.2f}σ")
        except tk.TclError:
            pass

    def on_central_multiplier_changed(self, *args):
        try:
            val = self.central_slider_val.get()
            self.lbl_central_formula.config(text=f"Central Threshold = μ - {val:.2f}σ")
        except tk.TclError:
            pass

    def log_message(self, message):
        self.txt_console.insert(tk.END, message + "\n")
        self.txt_console.see(tk.END)
        print(message)

    def update_slider_label(self, event=None):
        self.lbl_slider.config(text=f"Multiplier Factor: {self.slider_val.get():.2f}")

    def load_hdf5_file(self):
        file_path = filedialog.askopenfilename(filetypes=[("HDF5 Files", "*.h5;*.hdf5"), ("All Files", "*.*")])
        if not file_path:
            return

        self.h5_path = file_path
        self.lbl_filename.config(text=os.path.basename(file_path))

        for item in self.tree.get_children():
            self.tree.delete(item)
        self.datasets_dict.clear()

        try:
            with h5py.File(self.h5_path, "r") as f:
                def visitor(name, obj):
                    if isinstance(obj, h5py.Dataset):
                        shape_str = str(obj.shape)
                        dtype_str = str(obj.dtype)
                        node_id = self.tree.insert("", "end", text=name, values=(shape_str, dtype_str))
                        self.datasets_dict[node_id] = name

                f.visititems(visitor)
            self.log_message(f"Successfully unpacked structure for file: {os.path.basename(file_path)}")
        except Exception as e:
            messagebox.showerror("HDF5 Parsing Error", f"Unable to parse HDF5 elements:\n{str(e)}")

    def on_dataset_selected(self, event):
        selected_items = self.tree.selection()
        if not selected_items:
            return

        node_id = selected_items[0]
        if node_id not in self.datasets_dict:
            return

        dataset_path = self.datasets_dict[node_id]

        try:
            with h5py.File(self.h5_path, "r") as f:
                dset = f[dataset_path]
                shape = dset.shape

                self.ent_row_start.delete(0, tk.END)
                self.ent_row_start.insert(0, "1")
                self.ent_row_end.delete(0, tk.END)
                self.ent_row_end.insert(0, str(shape[0]))

                self.ent_col_start.delete(0, tk.END)
                self.ent_col_start.insert(0, "0")
                self.ent_col_end.delete(0, tk.END)
                self.ent_col_end.insert(0, str(shape[1] if len(shape) > 1 else 1))

                r_start = 1
                r_end = min(shape[0], 200003)

                if len(shape) == 1:
                    sample_data = dset[r_start:r_end]
                    target_display = "Array"
                else:
                    try:
                        c_start = int(self.ent_col_start.get())
                    except ValueError:
                        c_start = 0
                    sample_data = dset[r_start:r_end, c_start]
                    target_display = f"Col {c_start}"

                sample_data = sample_data[np.isfinite(sample_data)]

                if sample_data.size > 1:
                    sample_mean = np.mean(sample_data)
                    sample_std = np.std(sample_data)

                    self.lbl_pre_scan_stats.config(
                        text=f"Baseline ({target_display}): Mean={sample_mean:.5f} | Std Dev (σ)={sample_std:.5f}"
                    )
                else:
                    self.lbl_pre_scan_stats.config(text="Baseline: Matrix data vector sizing inadequate.")

                self.log_message(f"Focused Dataset: {dataset_path} | Native Matrix Base Shape: {shape}")
        except Exception as e:
            self.lbl_pre_scan_stats.config(text="Baseline: Error loading channel metrics.")
            self.log_message(f"Error reading shape or metrics: {str(e)}")

    def start_processing(self):
        selected_items = self.tree.selection()
        if not selected_items or (not self.h5_path):
            messagebox.showwarning("Selection Missing",
                                   "Please select an internal dataset path inside an active HDF5 file first!")
            return

        node_id = selected_items[0]
        self.active_dataset_name = self.datasets_dict[node_id]

        try:
            r_start = int(self.ent_row_start.get())
            r_end = int(self.ent_row_end.get())
            c_start = int(self.ent_col_start.get())
            c_end = int(self.ent_col_end.get())

            self.c_region_start = int(self.ent_central_start.get())
            self.c_region_end = int(self.ent_central_end.get())
        except ValueError:
            messagebox.showerror("Configuration Parameter Error",
                                 "Check bounds criteria. Please specify numerical integers.")
            return

        self.r_start, self.r_end = r_start, r_end
        self.columns_to_process = list(range(c_start, c_end))
        self.current_col_index = 0

        self.all_peak_counts = []
        self.peak_results = []
        self.doc_summary_table_data = []

        self.is_running = True
        self.btn_play.config(state=tk.DISABLED)
        self.btn_stop.config(state=tk.NORMAL)

        self.process_next_column_loop()

    def stop_processing(self):
        self.is_running = False
        self.btn_play.config(state=tk.NORMAL)
        self.btn_stop.config(state=tk.DISABLED)
        self.log_message("Sequence Paused by user.")

    def process_next_column_loop(self):
        if not self.is_running:
            return

        if self.current_col_index >= len(self.columns_to_process):
            self.finalize_document_generation()
            self.stop_processing()
            messagebox.showinfo("Task Successful",
                                "Finished scanning metrics across targeted bounds! Document compiled successfully.")
            return

        col_id = self.columns_to_process[self.current_col_index]
        self.analyze_single_column(col_id)
        self.current_col_index += 1
        self.root.after(100, self.process_next_column_loop)

    def analyze_single_column(self, i):
        try:
            with h5py.File(self.h5_path, "r") as f:
                dset = f[self.active_dataset_name]

                if len(dset.shape) == 1:
                    data = dset[self.r_start:self.r_end]
                    i = 0
                else:
                    data = dset[self.r_start:self.r_end, i]

            mean = np.mean(data)
            std = np.std(data)
            global_multiplier = self.slider_val.get()
            central_multiplier = self.central_slider_val.get()

            Global_Threshold = mean - (global_multiplier * std)
            Central_Threshold = mean - (central_multiplier * std)

            self.log_message(f"\n--- Processing Index Trace (Column): {i} ---")
            self.log_message(f"Signal Mean: {mean:.3f} | Std Dev: {std:.3f}")
            self.log_message(f"Global Threshold Limit: {Global_Threshold:.3f}")
            self.log_message(
                f"Central Threshold Limit: {Central_Threshold:.3f} (Span: {self.c_region_start} - {self.c_region_end})")

            peak_count = 0
            last_peak = -10000
            index = 0

            self.ax.clear()
            self.ax.plot(data, color="black", linewidth=1, label=f"Trace Column {i}")

            # LINE 1: Continuous Global Threshold across entire width
            self.ax.axhline(y=Global_Threshold, color='g', linestyle='--',
                            label=f'Global Threshold (μ - {global_multiplier:.1f}σ)')

            # LINE 2: Central threshold restricting its x-range exactly to user inputs
            c_start_idx = max(0, self.c_region_start)
            c_end_idx = min(len(data), self.c_region_end)
            if c_start_idx < c_end_idx:
                x_central = np.arange(c_start_idx, c_end_idx)
                y_central = np.full(c_end_idx - c_start_idx, Central_Threshold)
                self.ax.plot(x_central, y_central, color='m', linestyle='-', linewidth=2.5,
                             label=f'Central Threshold (μ - {central_multiplier:.1f}σ)')

            # Internal analytical processing threshold framework array
            threshold_array = np.full(len(data), Global_Threshold)
            if c_start_idx < c_end_idx:
                threshold_array[c_start_idx:c_end_idx] = Central_Threshold

            desktop_path = os.path.join(os.path.expanduser("~"), "Desktop")
            sanitized_dataset_name = self.active_dataset_name.replace("/", "_")
            h5_file_prefix = os.path.splitext(os.path.basename(self.h5_path))[0]

            peak_output_dir = os.path.join(desktop_path, f"Peaks_{h5_file_prefix}_{sanitized_dataset_name}")
            os.makedirs(peak_output_dir, exist_ok=True)

            while index < len(data):
                if data[index] >= threshold_array[index]:
                    index += 1
                    continue

                new_index = index + 6
                while new_index < len(data) and data[new_index] < threshold_array[new_index]:
                    new_index += 1

                start = index
                end = min(new_index + 1, len(data))

                x_arr = np.arange(start, end)
                y_arr = data[start:end]

                if len(y_arr) >= 4:
                    A0 = np.min(y_arr) - np.max(y_arr)
                    mu0 = x_arr[np.argmin(y_arr)]
                    sigma0 = max((end - start) / 4, 1)
                    C0 = np.max(y_arr)

                    try:
                        popt, _ = curve_fit(gaussian, x_arr, y_arr, p0=[A0, mu0, sigma0, C0])
                        peak_position = int(round(popt[1]))

                        # Keep lockout tight (200 indices) so peaks aren't missed
                        if peak_position - last_peak >= 200:
                            last_peak = peak_position
                            peak_count += 1

                            self.log_message(
                                f" -> Peak {peak_count} Fitted at Index Position: {peak_position} | Amp: {popt[0]:.2f}")

                            xfit = np.linspace(x_arr.min(), x_arr.max(), 300)
                            self.ax.plot(xfit, gaussian(xfit, *popt), 'r-', linewidth=2)
                            self.ax.plot(x_arr, y_arr, 'bo', markersize=2)

                            peak_filename = f"Peak_{h5_file_prefix}_{sanitized_dataset_name}_col{i}_p{peak_count}.txt"
                            peak_file_path = os.path.join(peak_output_dir, peak_filename)

                            with open(peak_file_path, "w") as pf:
                                pf.write(f"HDF5 Source File: {self.h5_path}\n")
                                pf.write(f"Dataset Array Structure: {self.active_dataset_name}\n")
                                pf.write(f"Trace Source Column Index: {i}\n")
                                pf.write(f"Peak Count Identifier: {peak_count}\n")
                                pf.write(f"Peak Data Matrix Coordinates (Absolute): Start={start}, End={end}\n")
                                pf.write("---------------------- Fitted Gaussian Variables ----------------------\n")
                                pf.write(f"Amplitude (A) : {popt[0]}\n")
                                pf.write(f"Mean Center (mu) : {popt[1]}\n")
                                pf.write(f"Sigma Variance : {popt[2]}\n")
                                pf.write(f"Baseline Offset (C) : {popt[3]}\n")
                                pf.write(f"Rounded Target Calculated Peak Center Position Index : {peak_position}\n")

                    except RuntimeError:
                        pass

                index = end

            self.log_message(f"Total dynamic peaks logged inside waveform slice = {peak_count}")
            self.all_peak_counts.append(peak_count)
            self.doc_summary_table_data.append((str(i), str(peak_count)))

            self.ax.set_xlabel("Sample Matrix Indexes")
            self.ax.set_ylabel("Signal Amplitude Units")
            self.ax.grid(True)
            self.ax.legend(loc="upper right")
            self.canvas.draw()

        except Exception as e:
            self.log_message(f"Critical execution parsing error on column trace index {i}: {str(e)}")

    def finalize_document_generation(self):
        """Compiles standard automated processing doc report data directly to user's local Desktop space."""
        try:
            desktop_path = os.path.join(os.path.expanduser("~"), "Desktop")
            h5_file_prefix = os.path.splitext(os.path.basename(self.h5_path))[0]
            sanitized_dataset_name = self.active_dataset_name.replace("/", "_")

            output_filename = f"Peak_Detection_Report_{h5_file_prefix}_{sanitized_dataset_name}.docx"
            output_file_path = os.path.join(desktop_path, output_filename)

            document = Document()
            document.add_heading(f"HDF5 Peak Analysis Summary: {h5_file_prefix}", level=1)
            document.add_paragraph(f"Source Internal Dataset Element Array Matrix Path: {self.active_dataset_name}")
            document.add_paragraph(f"Configured Row Limits Selected Range: From {self.r_start} To {self.r_end}")
            document.add_paragraph(
                f"Global Threshold Multiplier Coefficient Factor: {self.slider_val.get():.2f}")
            document.add_paragraph(
                f"Central Area Window Range: {self.c_region_start} to {self.c_region_end} | Multiplier Factor: {self.central_slider_val.get():.2f}")

            table = document.add_table(rows=1, cols=2)
            table.style = "Table Grid"
            hdr = table.rows[0].cells
            hdr[0].text = "Trace Index Number (Column)"
            hdr[1].text = "Calculated Isolated Peaks Count"

            for trace_idx, p_count in self.doc_summary_table_data:
                row_cells = table.add_row().cells
                row_cells[0].text = trace_idx
                row_cells[1].text = p_count

            if self.all_peak_counts:
                average_peaks = np.mean(self.all_peak_counts)
                document.add_heading("Run Analytics Metric Breakdown Summary", level=2)
                document.add_paragraph(
                    f"Average Peak Incidences structural average value across columns: {average_peaks:.2f}")

            document.save(output_file_path)
            self.log_message(
                f"\n[SAVED WORD REPORT] Successfully compiled and output file to path:\n -> {output_file_path}")
        except Exception as e:
            self.log_message(f"Failed parsing/compiling output document report data elements onto workspace: {str(e)}")


if __name__ == "__main__":
    root = tk.Tk()
    app = H5DataAnalyzerApp(root)


    def on_gui_close():
        app.is_running = False
        root.withdraw()
        root.quit()
        root.destroy()


    root.protocol("WM_DELETE_WINDOW", on_gui_close)
    root.mainloop()
